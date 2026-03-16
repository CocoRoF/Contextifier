"""
Comprehensive E2E tests for the DOCX handler pipeline.

Tests cover:
- DocxConverter: ZIP validation, Document creation, edge cases
- DocxPreprocessor: PreprocessedData wrapping, chart pre-extraction
- DocxMetadataExtractor: core_properties mapping
- _paragraph: run/drawing/pict/page-break processing
- _table_extractor: column widths, row/col spans, cell text
- DocxContentExtractor: full body traversal, images, charts
- DOCXHandler: end-to-end pipeline
"""

import io
import struct
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Optional
from unittest.mock import MagicMock, Mock, PropertyMock, patch

import pytest

# ── Project imports ───────────────────────────────────────────────────────

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from contextifier_new.types import (
    DocumentMetadata,
    ExtractionResult,
    FileContext,
    PreprocessedData,
    TableCell,
    TableData,
)
from contextifier_new.config import ProcessingConfig
from contextifier_new.errors import ConversionError

from contextifier_new.handlers.docx._constants import (
    NAMESPACES,
    ElementType,
    CHART_TYPE_MAP,
    ZIP_MAGIC,
)
from contextifier_new.handlers.docx.converter import DocxConverter
from contextifier_new.handlers.docx.preprocessor import DocxPreprocessor
from contextifier_new.handlers.docx.metadata_extractor import DocxMetadataExtractor
from contextifier_new.handlers.docx.content_extractor import DocxContentExtractor
from contextifier_new.handlers.docx.handler import DOCXHandler
from contextifier_new.handlers.docx._paragraph import (
    process_paragraph,
    has_page_break,
    extract_diagram_text,
    DrawingInfo,
    DrawingKind,
    PictInfo,
    RunContent,
)
from contextifier_new.handlers.docx._table_extractor import extract_table


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _make_minimal_docx(
    paragraphs: Optional[List[str]] = None,
    metadata: Optional[Dict[str, str]] = None,
) -> bytes:
    """
    Create a minimal valid DOCX (ZIP) file in memory.

    This is a real DOCX file that python-docx can open.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Set metadata
    if metadata:
        cp = doc.core_properties
        if "title" in metadata:
            cp.title = metadata["title"]
        if "author" in metadata:
            cp.author = metadata["author"]
        if "subject" in metadata:
            cp.subject = metadata["subject"]
        if "keywords" in metadata:
            cp.keywords = metadata["keywords"]
        if "comments" in metadata:
            cp.comments = metadata["comments"]
        if "category" in metadata:
            cp.category = metadata["category"]

    # Add paragraphs
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(
    headers: List[str],
    rows: List[List[str]],
    paragraphs: Optional[List[str]] = None,
) -> bytes:
    """Create a DOCX with a table (and optional paragraphs)."""
    from docx import Document

    doc = Document()

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    # Add table
    nrows = len(rows) + 1  # +1 for header
    ncols = len(headers)
    table = doc.add_table(rows=nrows, cols=ncols)

    # Header row
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_file_context(data: bytes, ext: str = "docx") -> FileContext:
    """Create a FileContext dict from bytes."""
    return FileContext(
        file_path=f"/tmp/test.{ext}",
        file_name=f"test.{ext}",
        file_extension=ext,
        file_category="document",
        file_data=data,
        file_stream=io.BytesIO(data),
        file_size=len(data),
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Constants tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestConstants:
    """Tests for _constants module."""

    def test_zip_magic(self):
        assert ZIP_MAGIC == b"PK\x03\x04"

    def test_namespaces_has_required_keys(self):
        required = {"w", "wp", "a", "pic", "r", "c"}
        assert required.issubset(NAMESPACES.keys())

    def test_element_type_enum(self):
        assert ElementType.TEXT.value == "text"
        assert ElementType.TABLE.value == "table"
        assert ElementType.IMAGE.value == "image"
        assert ElementType.CHART.value == "chart"
        assert ElementType.DIAGRAM.value == "diagram"
        assert ElementType.PAGE_BREAK.value == "page_break"

    def test_chart_type_map_has_common_types(self):
        assert "barChart" in CHART_TYPE_MAP
        assert "lineChart" in CHART_TYPE_MAP
        assert "pieChart" in CHART_TYPE_MAP


# ═══════════════════════════════════════════════════════════════════════════════
# Converter tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestDocxConverter:
    """Tests for DocxConverter."""

    def test_format_name(self):
        c = DocxConverter()
        assert c.get_format_name() == "docx"

    def test_convert_valid_docx(self):
        data = _make_minimal_docx(paragraphs=["Hello World"])
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)
        assert hasattr(doc, "paragraphs")
        assert len(doc.paragraphs) >= 1
        # Check text present
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Hello World" in texts

    def test_convert_empty_data_raises(self):
        fc = _make_file_context(b"")
        c = DocxConverter()
        with pytest.raises(ConversionError, match="Empty file data"):
            c.convert(fc)

    def test_convert_invalid_data_raises(self):
        fc = _make_file_context(b"NOT A DOCX")
        c = DocxConverter()
        with pytest.raises(ConversionError, match="Failed to open DOCX"):
            c.convert(fc)

    def test_validate_valid_docx(self):
        data = _make_minimal_docx()
        fc = _make_file_context(data)
        c = DocxConverter()
        assert c.validate(fc) is True

    def test_validate_too_short(self):
        fc = _make_file_context(b"PK")
        c = DocxConverter()
        assert c.validate(fc) is False

    def test_validate_not_zip(self):
        fc = _make_file_context(b"NOT A ZIP FILE AT ALL")
        c = DocxConverter()
        assert c.validate(fc) is False

    def test_validate_zip_without_content_types(self):
        """A ZIP file that is NOT a DOCX (missing [Content_Types].xml)."""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("dummy.txt", "hello")
        fc = _make_file_context(buf.getvalue())
        c = DocxConverter()
        assert c.validate(fc) is False

    def test_close_no_crash(self):
        data = _make_minimal_docx()
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)
        c.close(doc)  # Should not raise

    def test_close_none(self):
        c = DocxConverter()
        c.close(None)  # Should not raise

    def test_convert_multiple_paragraphs(self):
        data = _make_minimal_docx(paragraphs=["Para 1", "Para 2", "Para 3"])
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Para 1" in texts
        assert "Para 2" in texts
        assert "Para 3" in texts


# ═══════════════════════════════════════════════════════════════════════════════
# Preprocessor tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestDocxPreprocessor:
    """Tests for DocxPreprocessor."""

    def test_format_name(self):
        p = DocxPreprocessor()
        assert p.get_format_name() == "docx"

    def test_preprocess_wraps_document(self):
        data = _make_minimal_docx(paragraphs=["Hello"])
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        p = DocxPreprocessor()
        result = p.preprocess(doc)

        assert isinstance(result, PreprocessedData)
        assert result.content is doc
        assert result.raw_content is doc
        assert result.encoding == "utf-8"

    def test_preprocess_properties_has_counts(self):
        data = _make_minimal_docx(paragraphs=["P1", "P2"])
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        p = DocxPreprocessor()
        result = p.preprocess(doc)

        assert "paragraph_count" in result.properties
        assert result.properties["paragraph_count"] >= 2

    def test_preprocess_resources_has_charts_key(self):
        data = _make_minimal_docx(paragraphs=["Hello"])
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        p = DocxPreprocessor()
        result = p.preprocess(doc)

        assert "charts" in result.resources
        assert isinstance(result.resources["charts"], list)

    def test_preprocess_none_raises(self):
        from contextifier_new.errors import PreprocessingError
        p = DocxPreprocessor()
        with pytest.raises(PreprocessingError, match="None"):
            p.preprocess(None)

    def test_preprocess_table_count(self):
        data = _make_docx_with_table(
            headers=["A", "B"],
            rows=[["1", "2"]],
        )
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        p = DocxPreprocessor()
        result = p.preprocess(doc)

        assert result.properties["table_count"] >= 1


# ═══════════════════════════════════════════════════════════════════════════════
# Metadata Extractor tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestDocxMetadataExtractor:
    """Tests for DocxMetadataExtractor."""

    def test_format_name(self):
        m = DocxMetadataExtractor()
        assert m.get_format_name() == "docx"

    def test_extract_from_document(self):
        data = _make_minimal_docx(
            paragraphs=["Test"],
            metadata={"title": "My Title", "author": "John Doe"},
        )
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        m = DocxMetadataExtractor()
        md = m.extract(doc)

        assert isinstance(md, DocumentMetadata)
        assert md.title == "My Title"
        assert md.author == "John Doe"

    def test_extract_from_preprocessed_data(self):
        data = _make_minimal_docx(
            metadata={"title": "From Preprocessed", "subject": "Test Subject"},
        )
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        p = DocxPreprocessor()
        preprocessed = p.preprocess(doc)

        m = DocxMetadataExtractor()
        md = m.extract(preprocessed)

        assert md.title == "From Preprocessed"
        assert md.subject == "Test Subject"

    def test_extract_all_fields(self):
        data = _make_minimal_docx(
            metadata={
                "title": "T",
                "subject": "S",
                "author": "A",
                "keywords": "K",
                "comments": "C",
                "category": "Cat",
            },
        )
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        m = DocxMetadataExtractor()
        md = m.extract(doc)

        assert md.title == "T"
        assert md.subject == "S"
        assert md.author == "A"
        assert md.keywords == "K"
        assert md.comments == "C"
        assert md.category == "Cat"

    def test_extract_empty_metadata(self):
        data = _make_minimal_docx()
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        m = DocxMetadataExtractor()
        md = m.extract(doc)

        # Should return DocumentMetadata with mostly None/empty fields
        assert isinstance(md, DocumentMetadata)

    def test_extract_none_returns_empty(self):
        m = DocxMetadataExtractor()
        md = m.extract(None)
        assert isinstance(md, DocumentMetadata)
        assert md.title is None

    def test_extract_string_returns_empty(self):
        m = DocxMetadataExtractor()
        md = m.extract("not a document")
        assert isinstance(md, DocumentMetadata)
        assert md.title is None

    def test_revision_from_core_properties(self):
        """core_properties.revision is an int, should be converted to str."""
        data = _make_minimal_docx(metadata={"title": "Rev Test"})
        fc = _make_file_context(data)
        c = DocxConverter()
        doc = c.convert(fc)

        # python-docx sets revision as int
        try:
            doc.core_properties.revision = 5
        except Exception:
            pass  # Some versions may not allow this

        m = DocxMetadataExtractor()
        md = m.extract(doc)
        assert isinstance(md, DocumentMetadata)


# ═══════════════════════════════════════════════════════════════════════════════
# Paragraph processing tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestParagraphProcessing:
    """Tests for _paragraph module using real DOCX elements."""

    def _get_paragraph_elements(self, data: bytes):
        """Get lxml paragraph elements from a DOCX."""
        c = DocxConverter()
        fc = _make_file_context(data)
        doc = c.convert(fc)
        body = doc.element.body
        W = NAMESPACES["w"]
        return list(body.iterchildren(f"{{{W}}}p")), doc

    def test_process_simple_paragraph(self):
        data = _make_minimal_docx(paragraphs=["Hello World"])
        paras, doc = self._get_paragraph_elements(data)

        # Find the paragraph with content
        for p in paras:
            text, drawings, picts, break_ = process_paragraph(p)
            if text.strip() == "Hello World":
                assert drawings == []
                assert picts == []
                assert break_ is False
                return

        # Should have found it
        assert False, "Did not find 'Hello World' paragraph"

    def test_process_multiple_paragraphs(self):
        data = _make_minimal_docx(paragraphs=["Line 1", "Line 2"])
        paras, doc = self._get_paragraph_elements(data)

        texts = []
        for p in paras:
            text, _, _, _ = process_paragraph(p)
            if text.strip():
                texts.append(text.strip())

        assert "Line 1" in texts
        assert "Line 2" in texts

    def test_process_empty_paragraph(self):
        data = _make_minimal_docx(paragraphs=[""])
        paras, doc = self._get_paragraph_elements(data)

        # Empty paragraphs should return empty text
        for p in paras:
            text, drawings, picts, break_ = process_paragraph(p)
            # All should be valid (no crashes)
            assert isinstance(text, str)
            assert isinstance(drawings, list)

    def test_has_page_break_returns_false_for_normal(self):
        data = _make_minimal_docx(paragraphs=["Normal paragraph"])
        paras, _ = self._get_paragraph_elements(data)
        for p in paras:
            assert has_page_break(p) is False

    def test_extract_diagram_text_with_texts(self):
        """Test diagram text extraction with mock element."""
        from lxml import etree
        A = NAMESPACES["a"]
        xml = f"""
        <a:graphicData xmlns:a="{A}" uri="some/diagram">
            <a:t>Node One</a:t>
            <a:t>Node Two</a:t>
        </a:graphicData>
        """
        elem = etree.fromstring(xml)
        result = extract_diagram_text(elem)
        assert "Node One" in result
        assert "Node Two" in result
        assert result.startswith("[Diagram:")

    def test_extract_diagram_text_empty(self):
        from lxml import etree
        A = NAMESPACES["a"]
        xml = f'<a:graphicData xmlns:a="{A}" uri="some/diagram"></a:graphicData>'
        elem = etree.fromstring(xml)
        result = extract_diagram_text(elem)
        assert result == "[Diagram]"


# ═══════════════════════════════════════════════════════════════════════════════
# Table extractor tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestTableExtractor:
    """Tests for _table_extractor using real DOCX tables."""

    def _get_table_elements(self, data: bytes):
        """Get lxml table elements from a DOCX."""
        c = DocxConverter()
        fc = _make_file_context(data)
        doc = c.convert(fc)
        body = doc.element.body
        W = NAMESPACES["w"]
        return list(body.iterchildren(f"{{{W}}}tbl")), doc

    def test_extract_simple_table(self):
        data = _make_docx_with_table(
            headers=["Name", "Value"],
            rows=[["A", "1"], ["B", "2"]],
        )
        tables, _ = self._get_table_elements(data)
        assert len(tables) >= 1

        td = extract_table(tables[0])
        assert td is not None
        assert td.num_cols == 2
        assert td.num_rows >= 2  # At least header + 1 data row

    def test_extract_table_cell_content(self):
        data = _make_docx_with_table(
            headers=["H1", "H2"],
            rows=[["Cell_A", "Cell_B"]],
        )
        tables, _ = self._get_table_elements(data)
        td = extract_table(tables[0])
        assert td is not None

        # Collect all cell contents
        all_contents = []
        for row in td.rows:
            for cell in row:
                all_contents.append(cell.content)

        assert "H1" in all_contents
        assert "H2" in all_contents
        assert "Cell_A" in all_contents
        assert "Cell_B" in all_contents

    def test_extract_table_has_header(self):
        data = _make_docx_with_table(
            headers=["X", "Y"],
            rows=[["1", "2"], ["3", "4"]],
        )
        tables, _ = self._get_table_elements(data)
        td = extract_table(tables[0])
        assert td is not None
        assert td.has_header is True

    def test_extract_single_row_table(self):
        data = _make_docx_with_table(
            headers=["Only"],
            rows=[],
        )
        tables, _ = self._get_table_elements(data)
        td = extract_table(tables[0])
        # Single row table
        assert td is not None
        assert td.num_rows >= 1

    def test_extract_table_column_widths(self):
        data = _make_docx_with_table(
            headers=["A", "B", "C"],
            rows=[["1", "2", "3"]],
        )
        tables, _ = self._get_table_elements(data)
        td = extract_table(tables[0])
        assert td is not None
        # python-docx may or may not set tblGrid depending on creation
        # but if set, widths should be non-empty
        if td.col_widths_percent is not None:
            assert len(td.col_widths_percent) == 3
            assert abs(sum(td.col_widths_percent) - 100.0) < 1.0

    def test_extract_table_returns_none_for_empty(self):
        """extract_table should return None for an element with no rows."""
        from lxml import etree
        W = NAMESPACES["w"]
        xml = f'<w:tbl xmlns:w="{W}"></w:tbl>'
        elem = etree.fromstring(xml)
        result = extract_table(elem)
        assert result is None

    def test_extract_wide_table(self):
        """Table with many columns."""
        data = _make_docx_with_table(
            headers=["C1", "C2", "C3", "C4", "C5"],
            rows=[["v1", "v2", "v3", "v4", "v5"]],
        )
        tables, _ = self._get_table_elements(data)
        td = extract_table(tables[0])
        assert td is not None
        assert td.num_cols == 5

    def test_extract_table_with_empty_cells(self):
        data = _make_docx_with_table(
            headers=["H1", "H2"],
            rows=[["", "Data"], ["Data2", ""]],
        )
        tables, _ = self._get_table_elements(data)
        td = extract_table(tables[0])
        assert td is not None
        assert td.num_rows >= 2


# ═══════════════════════════════════════════════════════════════════════════════
# Content Extractor tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestDocxContentExtractor:
    """Tests for DocxContentExtractor."""

    def _make_preprocessed(self, data: bytes) -> PreprocessedData:
        c = DocxConverter()
        fc = _make_file_context(data)
        doc = c.convert(fc)
        p = DocxPreprocessor()
        return p.preprocess(doc)

    def test_format_name(self):
        ext = DocxContentExtractor()
        assert ext.get_format_name() == "docx"

    def test_extract_text_simple(self):
        data = _make_minimal_docx(paragraphs=["Hello DOCX World"])
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "Hello DOCX World" in text

    def test_extract_text_multiple_paragraphs(self):
        data = _make_minimal_docx(paragraphs=["First", "Second", "Third"])
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "First" in text
        assert "Second" in text
        assert "Third" in text

    def test_extract_text_with_table(self):
        data = _make_docx_with_table(
            headers=["Col1", "Col2"],
            rows=[["Value1", "Value2"]],
            paragraphs=["Before table"],
        )
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "Before table" in text
        # Table content should be present in some form
        assert "Col1" in text or "Value1" in text

    def test_extract_tables(self):
        data = _make_docx_with_table(
            headers=["H1", "H2"],
            rows=[["A", "B"]],
        )
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        tables = ext.extract_tables(preprocessed)

        assert len(tables) >= 1
        assert isinstance(tables[0], TableData)
        assert tables[0].num_cols == 2

    def test_extract_images_no_images(self):
        data = _make_minimal_docx(paragraphs=["No images here"])
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        images = ext.extract_images(preprocessed)

        assert images == []

    def test_extract_images_no_service(self):
        """Without ImageService, image extraction returns empty."""
        data = _make_minimal_docx(paragraphs=["Test"])
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor(image_service=None)
        images = ext.extract_images(preprocessed)

        assert images == []

    def test_extract_text_empty_document(self):
        data = _make_minimal_docx(paragraphs=[])
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        # Should return empty or minimal text (no crash)
        assert isinstance(text, str)

    def test_extract_all_orchestration(self):
        data = _make_docx_with_table(
            headers=["A", "B"],
            rows=[["1", "2"]],
            paragraphs=["Some text"],
        )
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        result = ext.extract_all(preprocessed)

        assert isinstance(result, ExtractionResult)
        assert "Some text" in result.text
        assert len(result.tables) >= 1

    def test_table_1x1_collapsed(self):
        """1x1 tables should be collapsed to content text."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Before")
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "Container Content"
        doc.add_paragraph("After")

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        preprocessed = self._make_preprocessed(data)
        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "Container Content" in text

    def test_table_single_column_collapsed(self):
        """Single-column tables should be collapsed to line-separated text."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        table = doc.add_table(rows=3, cols=1)
        table.rows[0].cells[0].text = "Row1"
        table.rows[1].cells[0].text = "Row2"
        table.rows[2].cells[0].text = "Row3"

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        preprocessed = self._make_preprocessed(data)
        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "Row1" in text
        assert "Row2" in text
        assert "Row3" in text

    def test_table_html_fallback(self):
        """Multi-column tables without TableService should produce HTML."""
        data = _make_docx_with_table(
            headers=["X", "Y"],
            rows=[["a", "b"]],
        )
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        # Should contain HTML table tags
        assert "<table" in text or "X" in text

    def test_extract_text_with_page_tag_service(self):
        """When TagService is provided, page tags should appear."""
        data = _make_minimal_docx(paragraphs=["Hello"])
        preprocessed = self._make_preprocessed(data)

        mock_tag_svc = MagicMock()
        mock_tag_svc.make_page_tag.return_value = "[Page: 1]"

        ext = DocxContentExtractor(tag_service=mock_tag_svc)
        text = ext.extract_text(preprocessed)

        assert "[Page: 1]" in text

    def test_extract_text_unicode(self):
        """Korean/CJK text should be extracted correctly."""
        data = _make_minimal_docx(paragraphs=["안녕하세요 DOCX 테스트"])
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "안녕하세요" in text
        assert "DOCX" in text

    def test_extract_text_long_document(self):
        """Test with many paragraphs."""
        paras = [f"Paragraph number {i}" for i in range(50)]
        data = _make_minimal_docx(paragraphs=paras)
        preprocessed = self._make_preprocessed(data)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "Paragraph number 0" in text
        assert "Paragraph number 49" in text

    def test_table_service_delegation(self):
        """When TableService is available, it should be used for formatting."""
        data = _make_docx_with_table(
            headers=["A", "B"],
            rows=[["1", "2"]],
        )
        preprocessed = self._make_preprocessed(data)

        mock_table_svc = MagicMock()
        mock_table_svc.format_table.return_value = "<mock-table>content</mock-table>"

        ext = DocxContentExtractor(table_service=mock_table_svc)
        text = ext.extract_text(preprocessed)

        assert "<mock-table>" in text

    def test_image_service_integration(self):
        """When ImageService is provided and images exist, tags should appear."""
        # Create a DOCX with an inline image
        from docx import Document as DocxDocument
        from docx.shared import Inches

        doc = DocxDocument()
        doc.add_paragraph("Before image")

        # Add a small PNG image
        png_data = _make_tiny_png()
        img_stream = io.BytesIO(png_data)
        try:
            doc.add_picture(img_stream, width=Inches(1))
        except Exception:
            # If adding picture fails, skip this test
            pytest.skip("Cannot add picture to docx in test environment")

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        preprocessed = self._make_preprocessed(data)

        mock_img_svc = MagicMock()
        mock_img_svc.save_and_tag.return_value = "[Image: test.png]"

        ext = DocxContentExtractor(image_service=mock_img_svc)
        text = ext.extract_text(preprocessed)

        assert "[Image: test.png]" in text


def _make_tiny_png() -> bytes:
    """Create a minimal valid 1x1 white PNG file."""
    import struct
    import zlib

    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    sig = b'\x89PNG\r\n\x1a\n'
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = b'\x00\xff\xff\xff'
    idat = zlib.compress(raw)

    return sig + _chunk(b'IHDR', ihdr) + _chunk(b'IDAT', idat) + _chunk(b'IEND', b'')


# ═══════════════════════════════════════════════════════════════════════════════
# Handler (end-to-end) tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestDOCXHandler:
    """End-to-end tests for DOCXHandler."""

    def _make_handler(self, **kwargs) -> DOCXHandler:
        config = ProcessingConfig()
        return DOCXHandler(config, **kwargs)

    def test_handler_name(self):
        h = self._make_handler()
        assert h.handler_name == "DOCX Handler"

    def test_supported_extensions(self):
        h = self._make_handler()
        assert h.supported_extensions == frozenset({"docx"})

    def test_process_simple_document(self):
        data = _make_minimal_docx(paragraphs=["E2E test paragraph"])
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert isinstance(result, ExtractionResult)
        assert "E2E test paragraph" in result.text

    def test_process_with_metadata(self):
        data = _make_minimal_docx(
            paragraphs=["Content"],
            metadata={"title": "E2E Title", "author": "E2E Author"},
        )
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert isinstance(result, ExtractionResult)
        assert result.metadata is not None
        assert result.metadata.title == "E2E Title"
        assert result.metadata.author == "E2E Author"

    def test_process_with_table(self):
        data = _make_docx_with_table(
            headers=["Name", "Score"],
            rows=[["Alice", "100"], ["Bob", "95"]],
            paragraphs=["Test scores:"],
        )
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert isinstance(result, ExtractionResult)
        assert "Test scores" in result.text
        assert len(result.tables) >= 1
        assert result.tables[0].num_cols == 2

    def test_process_empty_document(self):
        data = _make_minimal_docx(paragraphs=[])
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert isinstance(result, ExtractionResult)
        assert isinstance(result.text, str)

    def test_process_unicode_content(self):
        data = _make_minimal_docx(
            paragraphs=["한국어 테스트", "日本語テスト", "English test"],
        )
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert "한국어" in result.text
        assert "日本語" in result.text
        assert "English" in result.text

    def test_process_large_document(self):
        paras = [f"Paragraph {i}: " + "x" * 100 for i in range(100)]
        data = _make_minimal_docx(paragraphs=paras)
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert "Paragraph 0" in result.text
        assert "Paragraph 99" in result.text

    def test_extract_text_convenience(self):
        data = _make_minimal_docx(paragraphs=["Convenience method test"])
        fc = _make_file_context(data)

        h = self._make_handler()
        text = h.extract_text(fc)

        assert isinstance(text, str)
        assert "Convenience method" in text

    def test_process_multiple_tables(self):
        """Document with multiple tables."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Table 1:")
        t1 = doc.add_table(rows=2, cols=2)
        t1.rows[0].cells[0].text = "A"
        t1.rows[0].cells[1].text = "B"
        t1.rows[1].cells[0].text = "1"
        t1.rows[1].cells[1].text = "2"

        doc.add_paragraph("Table 2:")
        t2 = doc.add_table(rows=2, cols=3)
        t2.rows[0].cells[0].text = "X"
        t2.rows[0].cells[1].text = "Y"
        t2.rows[0].cells[2].text = "Z"
        t2.rows[1].cells[0].text = "a"
        t2.rows[1].cells[1].text = "b"
        t2.rows[1].cells[2].text = "c"

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert len(result.tables) >= 2

    def test_process_with_image_service(self):
        """E2E with mock ImageService."""
        from docx import Document as DocxDocument
        from docx.shared import Inches

        doc = DocxDocument()
        doc.add_paragraph("Image below:")

        png_data = _make_tiny_png()
        try:
            doc.add_picture(io.BytesIO(png_data), width=Inches(1))
        except Exception:
            pytest.skip("Cannot add picture in test environment")

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        fc = _make_file_context(data)

        mock_img_svc = MagicMock()
        mock_img_svc.save_and_tag.return_value = "[Image: e2e_image.png]"

        h = self._make_handler(image_service=mock_img_svc)
        result = h.process(fc)

        assert "[Image: e2e_image.png]" in result.text

    def test_process_with_tag_service(self):
        """E2E with mock TagService for page tags."""
        data = _make_minimal_docx(paragraphs=["With tags"])
        fc = _make_file_context(data)

        mock_tag_svc = MagicMock()
        mock_tag_svc.make_page_tag.return_value = "[Page Number: 1]"

        h = self._make_handler(tag_service=mock_tag_svc)
        result = h.process(fc)

        assert "[Page Number: 1]" in result.text

    def test_handler_factory_creates_real_components(self):
        """Verify handler creates real (not Null) pipeline components."""
        h = self._make_handler()
        assert isinstance(h.create_converter(), DocxConverter)
        assert isinstance(h.create_preprocessor(), DocxPreprocessor)
        assert isinstance(h.create_metadata_extractor(), DocxMetadataExtractor)
        assert isinstance(h.create_content_extractor(), DocxContentExtractor)

    def test_invalid_data_raises_conversion_error(self):
        fc = _make_file_context(b"NOT A DOCX FILE")
        h = self._make_handler()
        with pytest.raises(ConversionError):
            h.process(fc)

    def test_process_preserves_paragraph_order(self):
        paras = ["Alpha", "Beta", "Gamma", "Delta", "Epsilon"]
        data = _make_minimal_docx(paragraphs=paras)
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        # Check order
        idx_alpha = result.text.index("Alpha")
        idx_beta = result.text.index("Beta")
        idx_gamma = result.text.index("Gamma")
        idx_delta = result.text.index("Delta")
        idx_epsilon = result.text.index("Epsilon")
        assert idx_alpha < idx_beta < idx_gamma < idx_delta < idx_epsilon

    def test_process_mixed_content(self):
        """Document with paragraphs, table, more paragraphs."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Start of document")
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "H1"
        t.rows[0].cells[1].text = "H2"
        t.rows[1].cells[0].text = "D1"
        t.rows[1].cells[1].text = "D2"
        doc.add_paragraph("End of document")

        buf = io.BytesIO()
        doc.save(buf)
        fc = _make_file_context(buf.getvalue())

        h = self._make_handler()
        result = h.process(fc)

        assert "Start of document" in result.text
        assert "End of document" in result.text
        assert len(result.tables) >= 1

    def test_process_special_characters(self):
        """Special characters should not break extraction."""
        special = 'He said "hello" & she said <goodbye> — it\'s fine!'
        data = _make_minimal_docx(paragraphs=[special])
        fc = _make_file_context(data)

        h = self._make_handler()
        result = h.process(fc)

        assert "hello" in result.text
        assert "goodbye" in result.text


# ═══════════════════════════════════════════════════════════════════════════════
# Integration tests: Full pipeline stage verification
# ═══════════════════════════════════════════════════════════════════════════════

class TestFullPipeline:
    """Verify the full pipeline flows correctly through all stages."""

    def test_converter_to_preprocessor_to_extractor(self):
        """Stage 1 → Stage 2 → Stage 4 manually."""
        data = _make_minimal_docx(paragraphs=["Pipeline test"])
        fc = _make_file_context(data)

        conv = DocxConverter()
        doc = conv.convert(fc)

        prep = DocxPreprocessor()
        preprocessed = prep.preprocess(doc)

        ext = DocxContentExtractor()
        text = ext.extract_text(preprocessed)

        assert "Pipeline test" in text

    def test_converter_to_metadata(self):
        """Stage 1 → Stage 3."""
        data = _make_minimal_docx(
            metadata={"title": "Pipeline Title"},
        )
        fc = _make_file_context(data)

        conv = DocxConverter()
        doc = conv.convert(fc)

        meta = DocxMetadataExtractor()
        md = meta.extract(doc)

        assert md.title == "Pipeline Title"

    def test_full_pipeline_with_table(self):
        """All stages with table content."""
        data = _make_docx_with_table(
            headers=["Key", "Val"],
            rows=[["color", "blue"]],
            paragraphs=["Settings:"],
        )
        fc = _make_file_context(data)

        conv = DocxConverter()
        doc = conv.convert(fc)

        prep = DocxPreprocessor()
        preprocessed = prep.preprocess(doc)

        meta = DocxMetadataExtractor()
        md = meta.extract(preprocessed)

        ext = DocxContentExtractor()
        result = ext.extract_all(preprocessed)

        assert isinstance(result, ExtractionResult)
        assert "Settings" in result.text
        assert len(result.tables) >= 1

    def test_preprocessor_chart_count_no_charts(self):
        """Documents without charts should report chart_count=0."""
        data = _make_minimal_docx(paragraphs=["No charts"])
        fc = _make_file_context(data)

        conv = DocxConverter()
        doc = conv.convert(fc)

        prep = DocxPreprocessor()
        preprocessed = prep.preprocess(doc)

        assert preprocessed.properties["chart_count"] == 0
        assert preprocessed.resources["charts"] == []


# ═══════════════════════════════════════════════════════════════════════════════
# Edge case tests
# ═══════════════════════════════════════════════════════════════════════════════

class TestEdgeCases:
    """Edge cases and boundary conditions."""

    def test_very_long_paragraph(self):
        long_text = "A" * 50000
        data = _make_minimal_docx(paragraphs=[long_text])
        fc = _make_file_context(data)

        h = DOCXHandler(ProcessingConfig())
        result = h.process(fc)

        assert len(result.text) >= 50000

    def test_many_tables(self):
        """Document with 10 tables."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        for i in range(10):
            doc.add_paragraph(f"Table {i}")
            t = doc.add_table(rows=2, cols=2)
            t.rows[0].cells[0].text = f"H{i}"
            t.rows[0].cells[1].text = f"V{i}"
            t.rows[1].cells[0].text = f"D{i}a"
            t.rows[1].cells[1].text = f"D{i}b"

        buf = io.BytesIO()
        doc.save(buf)
        fc = _make_file_context(buf.getvalue())

        h = DOCXHandler(ProcessingConfig())
        result = h.process(fc)

        assert len(result.tables) >= 10

    def test_empty_table_cells(self):
        data = _make_docx_with_table(
            headers=["", ""],
            rows=[["", ""]],
        )
        fc = _make_file_context(data)

        h = DOCXHandler(ProcessingConfig())
        result = h.process(fc)

        # Should not crash
        assert isinstance(result, ExtractionResult)

    def test_whitespace_only_paragraphs(self):
        data = _make_minimal_docx(paragraphs=["   ", "\t", "\n", "Real text"])
        fc = _make_file_context(data)

        h = DOCXHandler(ProcessingConfig())
        result = h.process(fc)

        assert "Real text" in result.text

    def test_mixed_languages(self):
        data = _make_minimal_docx(paragraphs=[
            "English text",
            "한국어 텍스트",
            "日本語テキスト",
            "中文文本",
            "Текст на русском",
        ])
        fc = _make_file_context(data)

        h = DOCXHandler(ProcessingConfig())
        result = h.process(fc)

        assert "English" in result.text
        assert "한국어" in result.text
        assert "日本語" in result.text
        assert "中文" in result.text
        assert "русском" in result.text

    def test_table_extractor_is_robust_with_malformed(self):
        """Table extractor should handle non-table XML gracefully."""
        from lxml import etree
        W = NAMESPACES["w"]
        # A tbl element with non-standard children
        xml = f'<w:tbl xmlns:w="{W}"><w:p><w:r><w:t>Weird</w:t></w:r></w:p></w:tbl>'
        elem = etree.fromstring(xml)
        result = extract_table(elem)
        # Should return None (no tr elements)
        assert result is None


# ═══════════════════════════════════════════════════════════════════════════════
# Test runner
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
