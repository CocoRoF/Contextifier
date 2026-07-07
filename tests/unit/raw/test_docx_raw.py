# tests/unit/raw/test_docx_raw.py
"""DocxRawDocument — addressing parity with python-docx, run-preserving
edits (the edit2docs P0-3 fix), and the byte-preservation contract.

Every fixture document is built with python-docx so the parity claims
are checked against the library whose addressing edit2docs already
uses.
"""

from __future__ import annotations

import io
import struct
import zipfile
import zlib

import docx
import pytest
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn as docx_qn
from docx.oxml.parser import OxmlElement
from lxml import etree

from contextifier.raw import open_raw
from contextifier.raw.docx import DocxRawDocument
from contextifier.raw.xmlpart import NS, qn

W_DRAWING = qn("w:drawing")
W_HYPERLINK = qn("w:hyperlink")


# -- helpers ------------------------------------------------------------------


def _png_1x1() -> bytes:
    """A deterministic, valid 1x1 RGB PNG (no Pillow needed)."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00")
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", idat)
        + chunk(b"IEND", b"")
    )


def _doc_bytes(document) -> bytes:
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _raw(document) -> DocxRawDocument:
    return open_raw(_doc_bytes(document), extension="docx")


def _reopen(raw: DocxRawDocument):
    """Round-trip through save and reopen with python-docx."""
    return docx.Document(io.BytesIO(raw.to_bytes()))


def _parts(data: bytes) -> dict[str, bytes]:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return {n: z.read(n) for n in z.namelist() if not n.endswith("/")}


def _add_hyperlink(paragraph, text: str, url: str) -> str:
    """python-docx has no public hyperlink writer; build the oxml."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(docx_qn("r:id"), r_id)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)
    return r_id


def _set_cell_shading(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(docx_qn("w:val"), "clear")
    shd.set(docx_qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_shading_fill(cell) -> str | None:
    shd = cell._tc.find(f"{docx_qn('w:tcPr')}/{docx_qn('w:shd')}")
    return shd.get(docx_qn("w:fill")) if shd is not None else None


# -- 1. addressing parity -------------------------------------------------------


class TestAddressingParity:
    @pytest.fixture()
    def built(self):
        d = docx.Document()
        d.add_paragraph("Intro")
        d.add_paragraph("Section title", style="Heading 1")
        t1 = d.add_table(rows=2, cols=3)
        for r in range(2):
            for c in range(3):
                t1.cell(r, c).text = f"t1 r{r}c{c}"
        d.add_paragraph("Middle")
        t2 = d.add_table(rows=3, cols=3)
        t2.cell(0, 0).merge(t2.cell(0, 1))  # horizontal: gridSpan
        t2.cell(1, 2).merge(t2.cell(2, 2))  # vertical: vMerge
        seen: set[int] = set()
        for r in range(3):
            for c in range(3):
                cell = t2.cell(r, c)
                if id(cell._tc) not in seen:
                    seen.add(id(cell._tc))
                    cell.text = f"t2 r{r}c{c}"
        d.add_paragraph("End")
        data = _doc_bytes(d)
        return open_raw(data, extension="docx"), docx.Document(io.BytesIO(data))

    def test_paragraph_indices_and_texts_match_python_docx(self, built):
        raw, ref = built
        assert [p.text for p in raw.paragraphs] == [p.text for p in ref.paragraphs]
        assert [p.index for p in raw.paragraphs] == list(range(len(ref.paragraphs)))

    def test_paragraph_style_ids(self, built):
        raw, ref = built
        assert raw.paragraphs[0].style == "Normal"
        assert raw.paragraphs[1].style == ref.paragraphs[1].style.style_id
        assert raw.paragraphs[1].style == "Heading1"

    def test_table_dims(self, built):
        raw, ref = built
        assert len(raw.tables) == len(ref.tables) == 2
        for rt, ft in zip(raw.tables, ref.tables):
            assert rt.n_rows == len(ft.rows)
            assert rt.n_cols == len(ft.columns)

    def test_every_grid_cell_matches_including_merges(self, built):
        raw, ref = built
        for t, ft in enumerate(ref.tables):
            rt = raw.tables[t]
            for r in range(len(ft.rows)):
                for c in range(len(ft.columns)):
                    assert rt.cell(r, c).text == ft.rows[r].cells[c].text, (
                        f"table {t} cell ({r},{c})"
                    )

    def test_merged_positions_share_one_cell(self, built):
        raw, _ = built
        t2 = raw.tables[1]
        assert t2.cell(0, 0).element is t2.cell(0, 1).element  # gridSpan
        assert t2.cell(1, 2).element is t2.cell(2, 2).element  # vMerge

    def test_cell_out_of_range(self, built):
        raw, _ = built
        with pytest.raises(IndexError):
            raw.tables[0].cell(2, 0)
        with pytest.raises(IndexError):
            raw.tables[0].cell(0, 3)


# -- 2. run preservation (P0-3 fix) --------------------------------------------


class TestRunPreservingParagraphReplace:
    @pytest.fixture()
    def built(self):
        d = docx.Document()
        d.add_paragraph("plain first paragraph")
        p = d.add_paragraph()
        p.add_run("bold lead").bold = True
        p.add_run(" middle ")
        p.add_run().add_picture(io.BytesIO(_png_1x1()))
        p.add_run("italic tail").italic = True
        return d

    def test_replace_keeps_formatting_and_image(self, built):
        raw = _raw(built)
        raw.set_paragraph_text(1, "REPLACED")
        ref = _reopen(raw)

        p = ref.paragraphs[1]
        assert p.text == "REPLACED"
        # first run carried the text and kept its rPr
        assert p.runs[0].text == "REPLACED"
        assert p.runs[0].bold is True
        # image run survived, other pure-text runs are gone
        assert len(p._p.findall(f".//{docx_qn('w:drawing')}")) == 1
        assert len(p.runs) == 2

    def test_media_part_untouched(self, built):
        original = _parts(_doc_bytes(built))
        raw = _raw(built)
        raw.set_paragraph_text(1, "REPLACED")
        edited = _parts(raw.to_bytes())
        media = [n for n in original if n.startswith("word/media/")]
        assert media, "fixture should embed an image"
        for name in media:
            assert edited[name] == original[name]

    def test_whitespace_edges_get_space_preserve(self, built):
        raw = _raw(built)
        raw.set_paragraph_text(0, "  padded  ")
        assert _reopen(raw).paragraphs[0].text == "  padded  "

    def test_replace_on_run_free_paragraph_creates_run(self):
        d = docx.Document()
        d.add_paragraph()  # no runs at all
        raw = _raw(d)
        raw.set_paragraph_text(0, "fresh")
        assert _reopen(raw).paragraphs[0].text == "fresh"


# -- 3. cell image preservation ---------------------------------------------------


class TestCellSetTextPreservesImages:
    @pytest.fixture()
    def built(self):
        d = docx.Document()
        t = d.add_table(rows=1, cols=2)
        cell = t.cell(0, 0)
        cell.text = "before"
        cell.add_paragraph().add_run().add_picture(io.BytesIO(_png_1x1()))
        cell.add_paragraph("after")
        t.cell(0, 1).text = "sibling"
        return d

    def test_set_text_keeps_image_paragraph(self, built):
        raw = _raw(built)
        assert raw.tables[0].cell(0, 0).paragraph_count == 3
        raw.tables[0].cell(0, 0).set_text("X")
        ref = _reopen(raw)

        cell = ref.tables[0].cell(0, 0)
        texts = [p.text for p in cell.paragraphs]
        assert texts == ["X", "", ""]
        assert cell.text.strip() == "X"
        # the image paragraph kept its drawing, in position
        assert len(cell.paragraphs[1]._p.findall(f".//{docx_qn('w:drawing')}")) == 1
        # neighbours untouched
        assert ref.tables[0].cell(0, 1).text == "sibling"

    def test_media_part_untouched(self, built):
        original = _parts(_doc_bytes(built))
        raw = _raw(built)
        raw.tables[0].cell(0, 0).set_text("X")
        edited = _parts(raw.to_bytes())
        for name in (n for n in original if n.startswith("word/media/")):
            assert edited[name] == original[name]


# -- 4. hyperlink policy ------------------------------------------------------------


class TestHyperlinkPolicy:
    def _build(self):
        d = docx.Document()
        p = d.add_paragraph("Go to ")
        rid = _add_hyperlink(p, "click here", "https://example.com/")
        return d, rid

    def test_replace_keeps_hyperlink_element_and_rel(self):
        d, rid = self._build()
        raw = _raw(d)
        assert raw.paragraphs[0].text == "Go to click here"
        raw.set_paragraph_text(0, "REPLACED")

        p_el = raw.paragraphs[0].element
        links = p_el.findall(W_HYPERLINK)
        assert len(links) == 1 and links[0].get(qn("r:id")) == rid
        assert raw.paragraphs[0].text == "REPLACED"

        ref = _reopen(raw)
        assert ref.paragraphs[0].text == "REPLACED"
        assert rid in ref.part.rels  # relationship intact

    def test_hyperlink_carries_text_when_no_direct_run(self):
        d = docx.Document()
        p = d.add_paragraph()
        _add_hyperlink(p, "only link", "https://example.com/")
        raw = _raw(d)
        raw.set_paragraph_text(0, "NEW")
        p_el = raw.paragraphs[0].element
        assert len(p_el.findall(W_HYPERLINK)) == 1
        assert _reopen(raw).paragraphs[0].text == "NEW"

    def test_strip_empty_hyperlinks(self):
        d, rid = self._build()
        raw = _raw(d)
        raw.set_paragraph_text(0, "REPLACED")
        assert raw.strip_empty_hyperlinks(0) == 1
        assert raw.paragraphs[0].element.findall(W_HYPERLINK) == []

        ref = _reopen(raw)
        assert ref.paragraphs[0].text == "REPLACED"
        assert rid not in ref.part.rels  # orphaned rel cleaned up

    def test_strip_leaves_nonempty_hyperlinks(self):
        d, _ = self._build()
        raw = _raw(d)
        assert raw.strip_empty_hyperlinks(0) == 0
        assert len(raw.paragraphs[0].element.findall(W_HYPERLINK)) == 1


# -- 5. row editing ------------------------------------------------------------------


class TestRowEditing:
    @pytest.fixture()
    def built(self):
        d = docx.Document()
        t = d.add_table(rows=2, cols=3)
        for c in range(3):
            t.cell(0, c).text = f"head{c}"
            _set_cell_shading(t.cell(0, c), "FF0000")
            t.cell(1, c).text = f"data{c}"
        return d

    def test_insert_row_copies_template_formatting(self, built):
        raw = _raw(built)
        raw.tables[0].insert_row(1)  # template = row 0 (shaded)
        ref = _reopen(raw)

        t = ref.tables[0]
        assert (len(t.rows), len(t.columns)) == (3, 3)
        for c in range(3):
            assert t.cell(1, c).text == ""  # text cleared
            assert _cell_shading_fill(t.cell(1, c)) == "FF0000"  # props kept
            assert t.cell(0, c).text == f"head{c}"
            assert t.cell(2, c).text == f"data{c}"

    def test_insert_row_at_start_and_end(self, built):
        raw = _raw(built)
        raw.tables[0].insert_row(0)
        raw.tables[0].insert_row(raw.tables[0].n_rows)
        ref = _reopen(raw)
        assert len(ref.tables[0].rows) == 4
        assert ref.tables[0].cell(1, 0).text == "head0"
        assert ref.tables[0].cell(3, 0).text == ""

    def test_delete_row(self, built):
        raw = _raw(built)
        raw.tables[0].delete_row(0)
        ref = _reopen(raw)
        assert len(ref.tables[0].rows) == 1
        assert ref.tables[0].cell(0, 0).text == "data0"

    def test_row_index_bounds(self, built):
        raw = _raw(built)
        with pytest.raises(IndexError):
            raw.tables[0].insert_row(5)
        with pytest.raises(IndexError):
            raw.tables[0].delete_row(2)


# -- 6. paragraph insert/delete -----------------------------------------------------


class TestParagraphInsertDelete:
    @pytest.fixture()
    def built(self):
        d = docx.Document()
        for text in ("A", "B", "C"):
            d.add_paragraph(text)
        return d

    def test_insert_after_middle(self, built):
        raw = _raw(built)
        new = raw.insert_paragraph_after(1, "B2")
        assert new.index == 2
        assert [p.text for p in _reopen(raw).paragraphs] == ["A", "B", "B2", "C"]

    def test_insert_at_body_start(self, built):
        raw = _raw(built)
        new = raw.insert_paragraph_after(-1, "Start")
        assert new.index == 0
        assert [p.text for p in _reopen(raw).paragraphs] == ["Start", "A", "B", "C"]

    def test_insert_after_last_keeps_sectpr_last(self, built):
        raw = _raw(built)
        raw.insert_paragraph_after(2, "End", style="Heading1")
        data = raw.to_bytes()

        assert [p.text for p in docx.Document(io.BytesIO(data)).paragraphs] == [
            "A",
            "B",
            "C",
            "End",
        ]
        raw2 = open_raw(data, extension="docx")
        assert raw2.paragraphs[3].style == "Heading1"
        body = etree.fromstring(_parts(data)["word/document.xml"]).find("w:body", NS)
        assert body[-1].tag == qn("w:sectPr")  # sectPr still the last body child

    def test_delete_first_and_last(self, built):
        raw = _raw(built)
        raw.delete_paragraph(2)  # last paragraph, adjacent to sectPr
        raw.delete_paragraph(0)
        assert [p.text for p in _reopen(raw).paragraphs] == ["B"]

    def test_index_bounds(self, built):
        raw = _raw(built)
        with pytest.raises(IndexError):
            raw.set_paragraph_text(3, "x")
        with pytest.raises(IndexError):
            raw.delete_paragraph(-1)
        with pytest.raises(IndexError):
            raw.insert_paragraph_after(3, "x")


# -- 7. nested tables -----------------------------------------------------------------


class TestNestedTables:
    @pytest.fixture()
    def built(self):
        d = docx.Document()
        t = d.add_table(rows=2, cols=2)
        host = t.cell(0, 0)
        host.text = "host"
        nested = host.add_table(rows=2, cols=2)
        for r in range(2):
            for c in range(2):
                nested.cell(r, c).text = f"n{r}{c}"
        return d

    def test_nested_table_is_addressable(self, built):
        raw = _raw(built)
        assert len(raw.tables) == 1  # body-level only, matching python-docx
        nested = raw.tables[0].nested_tables(0, 0)
        assert len(nested) == 1
        assert (nested[0].n_rows, nested[0].n_cols) == (2, 2)
        assert nested[0].cell(1, 1).text == "n11"
        assert raw.tables[0].nested_tables(0, 1) == []

    def test_nested_cell_edit(self, built):
        raw = _raw(built)
        raw.tables[0].nested_tables(0, 0)[0].cell(0, 0).set_text("edited")
        ref = _reopen(raw)
        assert ref.tables[0].cell(0, 0).tables[0].cell(0, 0).text == "edited"

    def test_host_set_text_leaves_nested_table_intact(self, built):
        raw = _raw(built)
        raw.tables[0].cell(0, 0).set_text("HOST")
        ref = _reopen(raw)
        host = ref.tables[0].cell(0, 0)
        assert host.paragraphs[0].text == "HOST"
        assert len(host.tables) == 1
        assert host.tables[0].cell(0, 0).text == "n00"


# -- 8. byte preservation ---------------------------------------------------------------


class TestBytePreservation:
    def test_single_paragraph_edit_touches_only_document_xml(self):
        d = docx.Document()
        d.add_paragraph("hello")
        d.add_paragraph().add_run().add_picture(io.BytesIO(_png_1x1()))
        d.add_table(rows=1, cols=1).cell(0, 0).text = "cell"
        original = _parts(_doc_bytes(d))

        raw = open_raw(_doc_bytes(d), extension="docx")
        raw.set_paragraph_text(0, "changed")
        edited = _parts(raw.to_bytes())

        assert sorted(original) == sorted(edited)
        assert edited["word/document.xml"] != original["word/document.xml"]
        for name in original:
            if name != "word/document.xml":
                assert edited[name] == original[name], f"collateral change in {name}"


# -- 9. body order -------------------------------------------------------------------------


class TestBodyOrder:
    def test_document_order_of_paragraphs_and_tables(self):
        d = docx.Document()
        d.add_paragraph("p0")
        d.add_paragraph("p1")
        d.add_table(rows=1, cols=1)
        d.add_paragraph("p2")
        d.add_table(rows=1, cols=1)
        d.add_paragraph("p3")
        raw = _raw(d)
        assert raw.body_order() == [
            ("p", 0),
            ("p", 1),
            ("tbl", 0),
            ("p", 2),
            ("tbl", 1),
            ("p", 3),
        ]


# -- misc surface ---------------------------------------------------------------------------


class TestMiscSurface:
    def test_format_and_charts_empty(self):
        d = docx.Document()
        d.add_paragraph("x")
        raw = _raw(d)
        assert raw.format == "docx"
        assert raw.chart_part_names == []
        assert raw.charts == []

    def test_headers_and_footers_text(self):
        d = docx.Document()
        d.add_paragraph("body")
        section = d.sections[0]
        section.header.paragraphs[0].text = "the header"
        section.footer.paragraphs[0].text = "the footer"
        raw = _raw(d)
        assert list(raw.headers.values()) == ["the header"]
        assert list(raw.footers.values()) == ["the footer"]
        assert all(n.startswith("word/header") for n in raw.headers)
        assert all(n.startswith("word/footer") for n in raw.footers)

    def test_open_raw_dispatches_to_docx_model(self):
        d = docx.Document()
        raw = open_raw(_doc_bytes(d))  # no extension hint — sniffed
        assert isinstance(raw, DocxRawDocument)
